import io
import re
from collections import Counter
from typing import Optional

import streamlit as st
from docx import Document
from PyPDF2 import PdfReader


st.set_page_config(
    page_title="AI Resume Matcher",
    page_icon="🎯",
    layout="wide",
)


SKILLS = [
    "python",
    "java",
    "javascript",
    "typescript",
    "c++",
    "c#",
    "sql",
    "mysql",
    "postgresql",
    "oracle",
    "mongodb",
    "html",
    "css",
    "react",
    "angular",
    "vue",
    "node.js",
    "django",
    "flask",
    "fastapi",
    "streamlit",
    "machine learning",
    "deep learning",
    "artificial intelligence",
    "natural language processing",
    "nlp",
    "computer vision",
    "data analysis",
    "data analytics",
    "data science",
    "pandas",
    "numpy",
    "scikit-learn",
    "tensorflow",
    "pytorch",
    "keras",
    "matplotlib",
    "power bi",
    "tableau",
    "excel",
    "aws",
    "azure",
    "google cloud",
    "cloud computing",
    "docker",
    "kubernetes",
    "terraform",
    "linux",
    "git",
    "github",
    "jenkins",
    "ci/cd",
    "rest api",
    "api development",
    "microservices",
    "agile",
    "scrum",
    "cybersecurity",
    "networking",
    "technical support",
    "project management",
    "communication",
    "leadership",
    "problem solving",
    "collaboration",
]


STOP_WORDS = {
    "a",
    "about",
    "across",
    "after",
    "all",
    "also",
    "an",
    "and",
    "any",
    "are",
    "as",
    "at",
    "be",
    "been",
    "being",
    "but",
    "by",
    "can",
    "company",
    "do",
    "for",
    "from",
    "has",
    "have",
    "in",
    "into",
    "is",
    "it",
    "its",
    "job",
    "more",
    "most",
    "of",
    "on",
    "or",
    "our",
    "role",
    "that",
    "the",
    "their",
    "they",
    "this",
    "to",
    "using",
    "we",
    "will",
    "with",
    "work",
    "you",
    "your",
}


def extract_pdf_text(uploaded_file) -> str:
    """Extract selectable text from a PDF upload."""
    reader = PdfReader(uploaded_file)
    text_parts = []

    for page in reader.pages:
        page_text = page.extract_text()
        if page_text:
            text_parts.append(page_text)

    return "\n".join(text_parts)


def extract_docx_text(uploaded_file) -> str:
    """Extract paragraph and table text from a DOCX upload."""
    file_bytes = io.BytesIO(uploaded_file.getvalue())
    document = Document(file_bytes)

    text_parts = [paragraph.text for paragraph in document.paragraphs]

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                text_parts.append(cell.text)

    return "\n".join(text_parts)


def extract_resume_text(uploaded_file) -> str:
    """Choose the correct text extractor based on filename."""
    filename = uploaded_file.name.lower()

    if filename.endswith(".pdf"):
        return extract_pdf_text(uploaded_file)

    if filename.endswith(".docx"):
        return extract_docx_text(uploaded_file)

    return ""


def normalize_text(text: str) -> str:
    """Normalize punctuation, spacing, and capitalization."""
    text = text.lower()
    text = text.replace("–", "-").replace("—", "-")
    text = re.sub(r"[^a-z0-9+#./\-\s]", " ", text)
    text = re.sub(r"\s+", " ", text)
    return text.strip()


def tokenize(text: str) -> list[str]:
    """Return meaningful words while excluding common filler words."""
    words = re.findall(r"[a-zA-Z][a-zA-Z0-9+#./-]*", normalize_text(text))

    return [
        word
        for word in words
        if word not in STOP_WORDS and len(word) > 2
    ]


def find_skills(text: str) -> list[str]:
    """Find known technical and professional skills in text."""
    normalized = f" {normalize_text(text)} "
    found = []

    for skill in SKILLS:
        pattern = rf"(?<!\w){re.escape(skill)}(?!\w)"
        if re.search(pattern, normalized):
            found.append(skill)

    return sorted(set(found))


def keyword_overlap(
    resume_text: str,
    job_text: str,
) -> tuple[list[str], list[str]]:
    """Find useful matching and missing keywords."""
    resume_words = set(tokenize(resume_text))
    job_counts = Counter(tokenize(job_text))

    important_job_words = {
        word
        for word, count in job_counts.items()
        if count >= 1
    }

    matches = sorted(resume_words.intersection(important_job_words))
    missing = sorted(important_job_words - resume_words)

    return matches[:30], missing[:30]


def calculate_scores(
    resume_text: str,
    job_text: str,
    resume_skills: list[str],
    job_skills: list[str],
) -> dict[str, int]:
    """Calculate weighted ATS-style scores."""
    matching_keywords, _ = keyword_overlap(resume_text, job_text)

    job_words = set(tokenize(job_text))

    keyword_score = (
        len(matching_keywords) / max(len(job_words), 1)
    ) * 100

    if job_skills:
        skill_score = (
            len(set(resume_skills).intersection(job_skills))
            / len(job_skills)
        ) * 100
    else:
        skill_score = keyword_score

    resume_length = len(tokenize(resume_text))

    if 250 <= resume_length <= 900:
        formatting_score = 100
    elif 150 <= resume_length < 250 or 900 < resume_length <= 1200:
        formatting_score = 75
    else:
        formatting_score = 50

    overall_score = (
        skill_score * 0.60
        + keyword_score * 0.30
        + formatting_score * 0.10
    )

    return {
        "overall": min(round(overall_score), 100),
        "skills": min(round(skill_score), 100),
        "keywords": min(round(keyword_score), 100),
        "formatting": min(round(formatting_score), 100),
    }


def create_summary(resume_text: str, resume_skills: list[str]) -> str:
    """Create a short non-generative resume summary."""
    words = tokenize(resume_text)
    word_count = len(words)

    if resume_skills:
        skills_text = ", ".join(resume_skills[:8])
        return (
            f"The resume contains approximately {word_count} meaningful "
            f"words and highlights experience with {skills_text}."
        )

    return (
        f"The resume contains approximately {word_count} meaningful words. "
        "Add a clearly labeled Skills section to improve automated matching."
    )


def create_recommendations(
    scores: dict[str, int],
    missing_skills: list[str],
    resume_text: str,
) -> list[str]:
    """Build practical recommendations from the analysis."""
    recommendations = []

    if missing_skills:
        recommendations.append(
            "Add relevant missing skills only when you genuinely have "
            f"experience with them: {', '.join(missing_skills[:8])}."
        )

    if scores["keywords"] < 45:
        recommendations.append(
            "Use more exact terminology from the job description in your "
            "summary, skills, and experience bullets."
        )

    if scores["formatting"] < 80:
        recommendations.append(
            "Keep the resume focused, use standard headings, and aim for "
            "roughly one to two pages."
        )

    if not re.search(
        r"\b(increased|reduced|improved|created|developed|managed|led|built)\b",
        resume_text,
        re.IGNORECASE,
    ):
        recommendations.append(
            "Start achievement bullets with strong action verbs such as "
            "developed, built, improved, led, or automated."
        )

    if not re.search(r"\b\d+%|\$\d+|\b\d+\+?\b", resume_text):
        recommendations.append(
            "Add measurable results where possible, such as percentages, "
            "time saved, users supported, or projects completed."
        )

    if not recommendations:
        recommendations.append(
            "The resume is reasonably aligned. Review every suggested keyword "
            "and keep only truthful, job-relevant additions."
        )

    return recommendations


def score_message(score: int) -> str:
    if score >= 80:
        return "Strong match"
    if score >= 65:
        return "Good match"
    if score >= 45:
        return "Moderate match"
    return "Needs improvement"


def build_report(
    filename: str,
    scores: dict[str, int],
    matching_skills: list[str],
    missing_skills: list[str],
    matching_keywords: list[str],
    recommendations: list[str],
    summary: str,
) -> str:
    """Create a downloadable plain-text analysis report."""
    return f"""
AI RESUME MATCHER REPORT
========================

Resume: {filename}

OVERALL RESULT
--------------
Overall Match: {scores["overall"]}%
Skill Match: {scores["skills"]}%
Keyword Match: {scores["keywords"]}%
Resume Structure: {scores["formatting"]}%

RESUME SUMMARY
--------------
{summary}

MATCHING SKILLS
---------------
{", ".join(matching_skills) if matching_skills else "No recognized matching skills found."}

MISSING SKILLS
--------------
{", ".join(missing_skills) if missing_skills else "No recognized missing skills found."}

MATCHING KEYWORDS
-----------------
{", ".join(matching_keywords) if matching_keywords else "No meaningful matching keywords found."}

RECOMMENDATIONS
---------------
{chr(10).join(f"- {item}" for item in recommendations)}

Important: Add skills and claims only when they accurately represent your experience.
""".strip()


st.markdown(
    """
    <style>
        .main-title {
            font-size: 2.7rem;
            font-weight: 800;
            margin-bottom: 0;
        }

        .subtitle {
            color: #9aa0a6;
            margin-bottom: 2rem;
        }

        .skill-chip {
            display: inline-block;
            padding: 0.35rem 0.65rem;
            margin: 0.2rem;
            border-radius: 999px;
            background: rgba(49, 130, 206, 0.18);
            border: 1px solid rgba(49, 130, 206, 0.45);
        }
    </style>
    """,
    unsafe_allow_html=True,
)

st.markdown(
    '<p class="main-title">🎯 AI Resume Matcher</p>',
    unsafe_allow_html=True,
)
st.markdown(
    '<p class="subtitle">'
    "Analyze resume alignment, ATS keywords, skills, and improvement areas."
    "</p>",
    unsafe_allow_html=True,
)

uploaded_resume = st.file_uploader(
    "Upload your resume",
    type=["pdf", "docx"],
    help="Upload a text-based PDF or Microsoft Word document.",
)

job_text = st.text_area(
    "Paste the complete job description",
    height=280,
    placeholder=(
        "Paste the full responsibilities, qualifications, and required "
        "skills here..."
    ),
)

analyze = st.button(
    "Analyze Resume",
    type="primary",
    use_container_width=True,
)

if analyze:
    if uploaded_resume is None:
        st.warning("Upload a PDF or DOCX resume first.")
        st.stop()

    if not job_text.strip():
        st.warning("Paste the complete job description first.")
        st.stop()

    try:
        resume_text = extract_resume_text(uploaded_resume)
    except Exception as error:
        st.error(f"The resume could not be read: {error}")
        st.stop()

    if not resume_text.strip():
        st.error(
            "No selectable text was found. The resume may be a scanned image "
            "instead of a text-based PDF."
        )
        st.stop()

    resume_skills = find_skills(resume_text)
    job_skills = find_skills(job_text)

    matching_skills = sorted(
        set(resume_skills).intersection(job_skills)
    )
    missing_skills = sorted(
        set(job_skills) - set(resume_skills)
    )

    matching_keywords, missing_keywords = keyword_overlap(
        resume_text,
        job_text,
    )

    scores = calculate_scores(
        resume_text,
        job_text,
        resume_skills,
        job_skills,
    )

    summary = create_summary(resume_text, resume_skills)

    recommendations = create_recommendations(
        scores,
        missing_skills,
        resume_text,
    )

    report = build_report(
        uploaded_resume.name,
        scores,
        matching_skills,
        missing_skills,
        matching_keywords,
        recommendations,
        summary,
    )

    st.success("Resume analysis complete.")

    st.subheader("ATS Match Overview")

    col1, col2, col3, col4 = st.columns(4)

    col1.metric(
        "Overall Match",
        f'{scores["overall"]}%',
        score_message(scores["overall"]),
    )
    col2.metric("Skill Match", f'{scores["skills"]}%')
    col3.metric("Keyword Match", f'{scores["keywords"]}%')
    col4.metric("Resume Structure", f'{scores["formatting"]}%')

    st.progress(scores["overall"] / 100)

    st.subheader("Resume Summary")
    st.info(summary)

    left_column, right_column = st.columns(2)

    with left_column:
        st.subheader("✅ Matching Skills")

        if matching_skills:
            st.markdown(
                "".join(
                    f'<span class="skill-chip">{skill}</span>'
                    for skill in matching_skills
                ),
                unsafe_allow_html=True,
            )
        else:
            st.write("No recognized job skills matched.")

    with right_column:
        st.subheader("⚠️ Missing Skills")

        if missing_skills:
            for skill in missing_skills:
                st.write(f"- {skill}")
        else:
            st.write("No recognized required skills are missing.")

    st.subheader("Matching ATS Keywords")

    if matching_keywords:
        st.write(", ".join(matching_keywords))
    else:
        st.write("No meaningful matching keywords were found.")

    with st.expander("View additional missing keywords"):
        if missing_keywords:
            st.write(", ".join(missing_keywords))
        else:
            st.write("No additional missing keywords were found.")

    st.subheader("Resume Improvement Recommendations")

    for recommendation in recommendations:
        st.write(f"• {recommendation}")

    st.download_button(
        label="Download Analysis Report",
        data=report,
        file_name="resume_match_report.txt",
        mime="text/plain",
        use_container_width=True,
    )

    with st.expander("View extracted resume text"):
        st.text(resume_text[:10000])

st.caption(
    "This tool provides an estimated ATS-style comparison and does not "
    "guarantee an interview or hiring decision."
)